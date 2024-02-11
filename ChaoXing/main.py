import requests
import docx
from bs4 import BeautifulSoup

cookieStr = "k8sexam=1703153242.373.5971.194124; jrose=539FB6788D63FAD8275B9233E5AED18B.mooc-exam-1064682606-2cnkc; fid=2125; lv=1; _uid=243992615; uf=da0883eb5260151e1bbe1a72fec3b2321627c6c425fc23e12d3d88c91cf95d880223eccc8a3c281d36be8e4d53a01896913b662843f1f4ad6d92e371d7fdf644c848846cb427962dfd68be96b6183b1adbf915f2e4c9cd2ab0d4decd3c062efbb71187e27c611106; _d=1701759679872; UID=243992615; vc=3A520E4F01E5CB56B032A12736424D03; vc2=10811DC5036244393C12802E9715FA4D; vc3=J2Cdg5g8PmKdSYYcZAGxjeu3ebtW7xNGGLp94%2FIbag0xJlUPg0yTAbErEUfVgfTBy1zUKCB97aBfXr7jDHFBWaH7CNZeMJrdh8flJb%2FQrNkb3Tz8Jc%2BjbVnpVAbp5DWAar0Y6TFfO4lo7U5GzkWJ0JvRyG7m8%2BAzCDROpE9rjtI%3D67b577aad4afa3f8197c125b65be1c28; cx_p_token=24aa00c1eda366bd79c3723c32963ee0; p_auth_token=eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1aWQiOiIyNDM5OTI2MTUiLCJsb2dpblRpbWUiOjE3MDE3NTk2Nzk4NzQsImV4cCI6MTcwMjM2NDQ3OX0.Qnax8UJvMt563WiTXiypOrln2myrItUz52iBHQyvse8; xxtenc=70d35ab6c12b952936ad6d1255320f75; DSSTASH_LOG=C_38-UN_616-US_243992615-T_1701759679874; spaceFid=2125; spaceRoleId=""; k8s=1703153242.339.5971.929600; jrose=2E03DCC241C526ED29238747CCDEC4B8.mooc-2990459528-hxjnk; route=ac9a7739314fa6817cbac7e56032374b"

# cookieStr去空格
cookieStr = cookieStr.replace(" ", "")

# cookieStr按照分号切割
cookieList = cookieStr.split(";")

# 按照等号转为字典
cookieDict = {}
for cookie in cookieList:
    key, value = cookie.split("=")
    cookieDict[key] = value

# print(cookieDict)

# 设置ua
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.103 Safari/537.36",
    "Cookie": cookieStr
}
res = requests.get('https://mooc1.chaoxing.com/exam-ans/'
                   'exam/lookPaper?courseId=202913457'
                   '&classId=62616899&paperId=204414084&p=1&ut=s&cpi=271894216&examRelationId=2825105'
                   '&enc=e9b63ca98959d9f3fc00df8bdb6b434f&newMooc=true&openc=b3d3260750498fcbe2caf99050c70cd9',
                   headers=headers)
#
# print(res.text)
# 获取所有class=‘questionLi’
soup = BeautifulSoup(res.text, 'html.parser')
questionList = soup.find_all('div', class_='questionLi')

questionTitleList = []
answerList = []
# 取得文本
for question in questionList:
    # 取得题目
    questionTitleList.append(question.find('h3', class_='mark_name').get_text().replace('\n', '').replace('\t', ''))
    # 取得选项
    optionList = question.find_all('ul', class_='mark_letter')
    # 取得选项文本
    optionTextList = []
    for option in optionList:
        optionTextList.append(option.get_text().replace('<br>', '').replace('</br>', '').replace('&nbsp;', ' '))
    answerList.append(optionTextList)

# print(questionTitleList)
# print(answerList)

# max=3
# for i in range(1,len(questionTitleList)):
#     print(questionTitleList[i], end='\n')
#     for j in range(len(answerList[i])):
#         print(answerList[i][j], end='\n')


# 将上述结果写入word文件
document = docx.Document()
document.add_heading('题目', 0)
for i in range(1,len(questionTitleList)):
    document.add_heading(questionTitleList[i], 1)
    for j in range(len(answerList[i])):
        document.add_paragraph(answerList[i][j])
document.save('题目.docx')
